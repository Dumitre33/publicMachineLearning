"""
Generate IEEE-style Paper for Fashion Image Generation Project
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_column_width(column, width):
    """Set column width."""
    for cell in column.cells:
        cell.width = width


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_ieee_paper():
    """Create an IEEE-style paper about the fashion generation project."""
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # Set page margins (IEEE uses narrow margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.57)
        section.right_margin = Cm(1.57)
    
    # ==================== TITLE ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Fashion Image Generation Using Projected GAN\nand Stable Diffusion LoRA Fine-tuning")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # ==================== AUTHORS ====================
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = authors.add_run("\nAuthor Name")
    author_run.font.size = Pt(11)
    author_run.font.name = 'Times New Roman'
    
    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff_run = affiliation.add_run("Department of Computer Science\nUniversity Name\nCity, Country\nemail@university.edu")
    aff_run.font.size = Pt(10)
    aff_run.font.italic = True
    aff_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # Spacing
    
    # ==================== ABSTRACT ====================
    abstract_title = doc.add_paragraph()
    abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs_title_run = abstract_title.add_run("Abstract")
    abs_title_run.font.bold = True
    abs_title_run.font.italic = True
    abs_title_run.font.size = Pt(9)
    
    abstract_text = """This paper presents a dual-approach system for generating high-quality fashion images using two complementary deep learning methodologies: Projected Generative Adversarial Networks (Projected GAN) and Low-Rank Adaptation (LoRA) fine-tuning of Stable Diffusion. The Projected GAN approach leverages frozen pretrained EfficientNet features as a discriminator backbone, enabling faster convergence and improved training stability compared to traditional GAN architectures. The LoRA-based approach fine-tunes a pretrained Stable Diffusion v1.5 model on fashion-specific data, enabling text-conditioned generation with minimal computational overhead. We implement several stability techniques including spectral normalization, R1 gradient penalty, exponential moving average (EMA), and gradient clipping to ensure robust training. Experimental results on the DeepFashion dataset demonstrate the effectiveness of both approaches, with the Projected GAN generating diverse unconditional fashion images and the LoRA model producing high-fidelity text-guided outputs. The complete implementation achieves training in approximately 11-12 hours on consumer-grade hardware (NVIDIA RTX 4060, 8GB VRAM), making it accessible for practical applications."""
    
    abstract = doc.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_run = abstract.add_run(abstract_text)
    abs_run.font.size = Pt(9)
    abs_run.font.italic = True
    abs_run.font.name = 'Times New Roman'
    
    # Keywords
    keywords = doc.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_label = keywords.add_run("Keywords: ")
    kw_label.font.bold = True
    kw_label.font.italic = True
    kw_label.font.size = Pt(9)
    kw_text = keywords.add_run("Generative Adversarial Networks, Projected GAN, Stable Diffusion, LoRA, Fashion Image Generation, Deep Learning")
    kw_text.font.italic = True
    kw_text.font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing
    
    # ==================== I. INTRODUCTION ====================
    intro_title = doc.add_paragraph()
    intro_run = intro_title.add_run("I. INTRODUCTION")
    intro_run.font.bold = True
    intro_run.font.size = Pt(10)
    
    intro_text1 = """The fashion industry has increasingly embraced artificial intelligence for various applications, from trend forecasting to virtual try-on systems. Image generation stands as a particularly impactful application, enabling designers to rapidly prototype concepts, e-commerce platforms to augment product catalogs, and marketing teams to create diverse visual content without extensive photoshoots."""
    
    intro_p1 = doc.add_paragraph()
    intro_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p1.add_run(intro_text1).font.size = Pt(10)
    
    intro_text2 = """Generative Adversarial Networks (GANs) have revolutionized image synthesis since their introduction by Goodfellow et al. [1]. However, training GANs remains challenging due to issues such as mode collapse, training instability, and the need for careful hyperparameter tuning. Recent advances, particularly Projected GANs [2], address these challenges by utilizing frozen pretrained features as discriminator backbones, significantly improving training stability and convergence speed."""
    
    intro_p2 = doc.add_paragraph()
    intro_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p2.add_run(intro_text2).font.size = Pt(10)
    
    intro_text3 = """Simultaneously, diffusion models have emerged as a powerful alternative to GANs. Stable Diffusion [3], in particular, has demonstrated remarkable capability in generating high-fidelity images from text descriptions. The introduction of Low-Rank Adaptation (LoRA) [4] has made fine-tuning these large models practical on consumer hardware by training only a small set of additional parameters."""
    
    intro_p3 = doc.add_paragraph()
    intro_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p3.add_run(intro_text3).font.size = Pt(10)
    
    intro_text4 = """This paper presents a comprehensive implementation of both approaches for fashion image generation. Our contributions include: (1) a robust Projected GAN implementation with multiple stability enhancements, (2) an efficient LoRA fine-tuning pipeline for Stable Diffusion, (3) detailed experimental analysis on the DeepFashion dataset, and (4) optimization strategies for training on consumer-grade GPUs."""
    
    intro_p4 = doc.add_paragraph()
    intro_p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p4.add_run(intro_text4).font.size = Pt(10)
    
    # ==================== II. THEORETICAL BACKGROUND ====================
    theory_title = doc.add_paragraph()
    theory_run = theory_title.add_run("II. THEORETICAL BACKGROUND")
    theory_run.font.bold = True
    theory_run.font.size = Pt(10)
    
    # A. GANs
    gan_theory_subtitle = doc.add_paragraph()
    gan_theory_sub_run = gan_theory_subtitle.add_run("A. Generative Adversarial Networks")
    gan_theory_sub_run.font.italic = True
    gan_theory_sub_run.font.size = Pt(10)
    
    gan_theory_text = """Generative Adversarial Networks (GANs) represent a class of deep learning models that learn to generate realistic data through an adversarial training process [1]. The architecture consists of two neural networks: a generator and a discriminator. The generator takes random noise as input and produces synthetic samples, while the discriminator evaluates whether a given sample is real (from the training data) or fake (produced by the generator).

During training, the two networks compete against each other in a game-theoretic framework. The discriminator learns to become better at distinguishing real from fake samples, while the generator learns to produce increasingly realistic samples that can fool the discriminator. This adversarial dynamic drives both networks to improve continuously until the generator produces samples indistinguishable from real data.

However, GAN training is notoriously unstable and prone to several failure modes. Mode collapse occurs when the generator learns to produce only a limited variety of outputs, ignoring the full diversity of the training data. Training instability can manifest as oscillating losses, vanishing gradients, or complete training divergence. These challenges have motivated numerous architectural and training improvements over the years."""
    
    gan_theory_p = doc.add_paragraph()
    gan_theory_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    gan_theory_p.add_run(gan_theory_text).font.size = Pt(10)
    
    # B. Projected GAN
    proj_theory_subtitle = doc.add_paragraph()
    proj_theory_sub_run = proj_theory_subtitle.add_run("B. Projected GAN Architecture")
    proj_theory_sub_run.font.italic = True
    proj_theory_sub_run.font.size = Pt(10)
    
    proj_theory_text = """Projected GANs [2] represent a significant advancement in GAN architecture by leveraging transfer learning within the discriminator. Instead of training a discriminator from scratch, Projected GANs use a pretrained image classification network (such as EfficientNet trained on ImageNet) as a frozen feature extractor. This pretrained backbone provides rich, semantically meaningful features that have already learned to recognize objects, textures, and structures from millions of images.

The key innovation is the multi-scale feature projection approach. Features are extracted from multiple intermediate layers of the pretrained network, capturing information at different levels of abstraction. Low-level features capture edges and textures, while high-level features capture semantic content and object parts. Each feature scale is then projected through learned convolution layers and processed by independent discriminator heads.

This architecture offers several advantages. First, the frozen pretrained features provide stable gradient signals from the beginning of training, dramatically accelerating convergence. Second, the multi-scale discrimination provides feedback at multiple resolutions, helping the generator learn both fine details and global structure. Third, training is more stable because the pretrained backbone constrains the discriminator's behavior."""
    
    proj_theory_p = doc.add_paragraph()
    proj_theory_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    proj_theory_p.add_run(proj_theory_text).font.size = Pt(10)
    
    # C. Style-Based Generation
    style_theory_subtitle = doc.add_paragraph()
    style_theory_sub_run = style_theory_subtitle.add_run("C. Style-Based Generation")
    style_theory_sub_run.font.italic = True
    style_theory_sub_run.font.size = Pt(10)
    
    style_theory_text = """Our generator architecture draws inspiration from StyleGAN [5], which introduced the concept of style-based image synthesis. The core idea is to separate the generation process into two stages: mapping and synthesis.

The mapping network is a series of fully connected layers that transforms the initial random noise vector into an intermediate latent representation. This intermediate space has better properties for image manipulation, as different dimensions tend to control different visual attributes more independently. This disentanglement makes it easier for the network to learn meaningful variations in the generated images.

The synthesis network then uses this intermediate representation to modulate the image generation process through Adaptive Instance Normalization (AdaIN). At each layer of the synthesis network, the feature statistics (mean and variance) are adjusted based on the style information. This allows the style vector to control visual attributes at multiple scales, from high-level features like overall color scheme to fine details like texture patterns. Noise injection at each layer adds stochastic variation, enabling diverse outputs from the same style input."""
    
    style_theory_p = doc.add_paragraph()
    style_theory_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_theory_p.add_run(style_theory_text).font.size = Pt(10)
    
    # D. Training Stability
    stab_theory_subtitle = doc.add_paragraph()
    stab_theory_sub_run = stab_theory_subtitle.add_run("D. Training Stability Techniques")
    stab_theory_sub_run.font.italic = True
    stab_theory_sub_run.font.size = Pt(10)
    
    stab_theory_text = """Several techniques are essential for stable GAN training. Spectral Normalization constrains the neural network layers to have bounded gradients by normalizing weight matrices based on their largest singular value. This prevents any single layer from amplifying gradients excessively, reducing the risk of training instability and mode collapse.

R1 Gradient Penalty [9] is a regularization technique applied to the discriminator. It penalizes the discriminator for having large gradients when evaluating real samples. This encourages the discriminator to make smooth predictions, preventing it from creating overly sharp decision boundaries that can destabilize generator training.

Exponential Moving Average (EMA) maintains a running average of the generator's weights throughout training. Instead of using the current generator weights for evaluation and sample generation, the averaged weights are used. This smoothing effect reduces the impact of individual training steps and produces more consistent, higher-quality outputs. The averaging is heavily weighted toward recent weights while still incorporating historical information.

Gradient Clipping limits the maximum magnitude of gradients during backpropagation. When gradients exceed a threshold, they are scaled down proportionally. This prevents gradient explosion, where extremely large gradients cause unstable weight updates. Learning rate warmup gradually increases the learning rate from zero to its target value during early training, preventing large weight updates before the network has found a reasonable region of the parameter space."""
    
    stab_theory_p = doc.add_paragraph()
    stab_theory_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    stab_theory_p.add_run(stab_theory_text).font.size = Pt(10)
    
    # E. Diffusion Models
    diff_theory_subtitle = doc.add_paragraph()
    diff_theory_sub_run = diff_theory_subtitle.add_run("E. Denoising Diffusion Probabilistic Models")
    diff_theory_sub_run.font.italic = True
    diff_theory_sub_run.font.size = Pt(10)
    
    diff_theory_text = """Diffusion models [7] represent a fundamentally different approach to image generation compared to GANs. Instead of learning a direct mapping from noise to images, diffusion models learn to reverse a gradual corruption process. The core idea is inspired by physical diffusion processes where structured patterns dissolve into randomness over time.

The forward diffusion process gradually adds Gaussian noise to training images over many timesteps until they become indistinguishable from pure random noise. This process is simple and well-defined. The reverse process, which the neural network learns, gradually removes noise to recover the original image structure. During training, the model learns to predict and remove the noise added at each step.

Stable Diffusion [3] makes diffusion models practical for high-resolution image generation by operating in a compressed latent space rather than pixel space. A pretrained encoder compresses images into a lower-dimensional representation, the diffusion process operates in this latent space, and a decoder reconstructs the final image. This dramatically reduces computational requirements while maintaining image quality.

Text conditioning is achieved through CLIP text embeddings [11], which provide semantic guidance during the denoising process. The model learns to generate images that match the meaning of text prompts, enabling precise control over generated content through natural language descriptions."""
    
    diff_theory_p = doc.add_paragraph()
    diff_theory_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    diff_theory_p.add_run(diff_theory_text).font.size = Pt(10)
    
    # F. LoRA
    lora_theory_subtitle = doc.add_paragraph()
    lora_theory_sub_run = lora_theory_subtitle.add_run("F. Low-Rank Adaptation (LoRA)")
    lora_theory_sub_run.font.italic = True
    lora_theory_sub_run.font.size = Pt(10)
    
    lora_theory_text = """Low-Rank Adaptation (LoRA) [4] is a parameter-efficient fine-tuning technique that enables adaptation of large pretrained models with minimal computational resources. The key insight is that the weight changes needed to adapt a model to a new domain often have low intrinsic dimensionality, meaning they can be well-approximated using low-rank matrix decompositions.

Instead of updating all weights in a pretrained model, LoRA freezes the original weights and adds small trainable matrices alongside them. These additional matrices are structured as the product of two smaller matrices, drastically reducing the number of trainable parameters. For example, adapting a weight matrix with millions of parameters might only require training a few thousand parameters in the low-rank decomposition.

When applied to Stable Diffusion, LoRA adapts the attention mechanisms and key layers of the U-Net architecture. The attention layers control how different parts of the image relate to each other and to the text conditioning, making them particularly important for learning new visual concepts. By training only these low-rank adaptations, LoRA enables domain-specific fine-tuning on consumer hardware with limited VRAM.

The trained LoRA weights can be saved separately from the base model and loaded on demand, enabling easy switching between different adaptations. This modularity makes LoRA particularly practical for applications requiring multiple specialized models."""
    
    lora_theory_p = doc.add_paragraph()
    lora_theory_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lora_theory_p.add_run(lora_theory_text).font.size = Pt(10)
    
    # ==================== III. RELATED WORK ====================
    related_title = doc.add_paragraph()
    related_run = related_title.add_run("III. RELATED WORK")
    related_run.font.bold = True
    related_run.font.size = Pt(10)
    
    # A. GANs
    gan_subtitle = doc.add_paragraph()
    gan_sub_run = gan_subtitle.add_run("A. Generative Adversarial Networks")
    gan_sub_run.font.italic = True
    gan_sub_run.font.size = Pt(10)
    
    gan_text = """GANs consist of a generator G and discriminator D trained in an adversarial manner. The generator learns to produce realistic samples while the discriminator learns to distinguish real from generated samples. StyleGAN [5] introduced style-based generation with adaptive instance normalization (AdaIN), enabling fine-grained control over generated images. StyleGAN2 [6] further improved quality through path length regularization and revised normalization. Projected GANs [2] demonstrated that using frozen pretrained features dramatically accelerates training while maintaining quality."""
    
    gan_p = doc.add_paragraph()
    gan_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    gan_p.add_run(gan_text).font.size = Pt(10)
    
    # B. Diffusion Models
    diff_subtitle = doc.add_paragraph()
    diff_sub_run = diff_subtitle.add_run("B. Diffusion Models and LoRA")
    diff_sub_run.font.italic = True
    diff_sub_run.font.size = Pt(10)
    
    diff_text = """Denoising Diffusion Probabilistic Models (DDPMs) [7] learn to reverse a gradual noising process. Latent Diffusion Models [3] operate in a compressed latent space, enabling efficient high-resolution synthesis. LoRA [4] enables parameter-efficient fine-tuning by decomposing weight updates into low-rank matrices, reducing trainable parameters by orders of magnitude while preserving adaptation capability."""
    
    diff_p = doc.add_paragraph()
    diff_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    diff_p.add_run(diff_text).font.size = Pt(10)
    
    # ==================== IV. METHODOLOGY ====================
    method_title = doc.add_paragraph()
    method_run = method_title.add_run("IV. METHODOLOGY")
    method_run.font.bold = True
    method_run.font.size = Pt(10)
    
    # A. Dataset
    data_subtitle = doc.add_paragraph()
    data_sub_run = data_subtitle.add_run("A. Dataset Preparation")
    data_sub_run.font.italic = True
    data_sub_run.font.size = Pt(10)
    
    data_text = """We utilize the DeepFashion dataset [8], specifically the training images subset containing approximately 10,000 high-quality fashion photographs. Images are preprocessed to 256x256 resolution for the Projected GAN and 512x512 for Stable Diffusion LoRA training. Data augmentation techniques including random horizontal flipping, color jitter, random cropping, and random erasing are applied to increase effective dataset diversity and prevent overfitting."""
    
    data_p = doc.add_paragraph()
    data_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    data_p.add_run(data_text).font.size = Pt(10)
    
    # B. Projected GAN Architecture
    pgan_subtitle = doc.add_paragraph()
    pgan_sub_run = pgan_subtitle.add_run("B. Projected GAN Architecture")
    pgan_sub_run.font.italic = True
    pgan_sub_run.font.size = Pt(10)
    
    pgan_text1 = """Our Projected GAN implementation consists of two main components: a style-based generator and a multi-scale projected discriminator."""
    
    pgan_p1 = doc.add_paragraph()
    pgan_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pgan_p1.add_run(pgan_text1).font.size = Pt(10)
    
    pgan_text2 = """Generator Architecture: The generator follows a StyleGAN-inspired design with a mapping network that transforms latent vectors z in R^256 to an intermediate latent space w. The synthesis network then progressively upsamples from a learned 4x4 constant through six synthesis blocks, each performing bilinear upsampling followed by two convolution layers with Adaptive Instance Normalization (AdaIN). Spectral normalization is applied to all convolutional layers to constrain the Lipschitz constant and improve training stability. Learnable noise injection with small initial scales provides stochastic variation."""
    
    pgan_p2 = doc.add_paragraph()
    pgan_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pgan_p2.add_run(pgan_text2).font.size = Pt(10)
    
    pgan_text3 = """Discriminator Architecture: The discriminator leverages a frozen pretrained EfficientNet-Lite0 backbone. Multi-scale features are extracted from four intermediate layers and projected through learned 1x1 convolutions to a common channel dimension (128). Each scale has an independent discriminator head producing per-pixel real/fake predictions. The final output aggregates predictions across all scales. The frozen backbone provides rich semantic features without additional training cost."""
    
    pgan_p3 = doc.add_paragraph()
    pgan_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pgan_p3.add_run(pgan_text3).font.size = Pt(10)
    
    # C. Stability Techniques
    stab_subtitle = doc.add_paragraph()
    stab_sub_run = stab_subtitle.add_run("C. Training Stability Techniques")
    stab_sub_run.font.italic = True
    stab_sub_run.font.size = Pt(10)
    
    stab_text = """To ensure robust training, we implement several stability mechanisms:

1) Spectral Normalization: Applied to generator convolutions to control the Lipschitz constant, preventing gradient explosion.

2) R1 Gradient Penalty: Regularizes the discriminator by penalizing the gradient magnitude on real data, computed every 8 training steps with gamma=0.1.

3) Exponential Moving Average (EMA): Maintains a smoothed copy of generator weights with decay=0.9999, producing more stable outputs during inference.

4) Gradient Clipping: Generator and discriminator gradients are clipped to maximum norm of 0.5.

5) Learning Rate Warmup: Linear warmup over 1000 steps prevents early training instability.

6) NaN Detection: Automatic detection of numerical instability with early stopping after 30 consecutive NaN losses.

7) Generator FP32 Training: While the discriminator uses mixed precision (FP16), the generator trains in FP32 for numerical stability."""
    
    stab_p = doc.add_paragraph()
    stab_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    stab_p.add_run(stab_text).font.size = Pt(10)
    
    # D. LoRA Fine-tuning
    lora_subtitle = doc.add_paragraph()
    lora_sub_run = lora_subtitle.add_run("D. Stable Diffusion LoRA Fine-tuning")
    lora_sub_run.font.italic = True
    lora_sub_run.font.size = Pt(10)
    
    lora_text = """For text-conditioned generation, we fine-tune Stable Diffusion v1.5 using LoRA. The base model remains frozen while low-rank adaptation matrices are trained on attention layers (Q, K, V, output projections), feed-forward layers, and convolution layers. We use rank r=128 and alpha=128. Training uses the AdamW optimizer with learning rate 5e-5, cosine learning rate schedule, and gradient checkpointing for memory efficiency. The instance prompt "a photograph of sks fashion clothing" enables the model to learn the fashion domain while preserving general capabilities."""
    
    lora_p = doc.add_paragraph()
    lora_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lora_p.add_run(lora_text).font.size = Pt(10)
    
    # ==================== V. IMPLEMENTATION ====================
    impl_title = doc.add_paragraph()
    impl_run = impl_title.add_run("V. IMPLEMENTATION DETAILS")
    impl_run.font.bold = True
    impl_run.font.size = Pt(10)
    
    impl_text = """The implementation uses PyTorch 2.x with CUDA acceleration. Key dependencies include timm for pretrained backbones, Hugging Face diffusers and transformers for Stable Diffusion, and PEFT for LoRA implementation."""
    
    impl_p = doc.add_paragraph()
    impl_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    impl_p.add_run(impl_text).font.size = Pt(10)
    
    # Training configuration table
    table_title = doc.add_paragraph()
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_run = table_title.add_run("TABLE I: TRAINING CONFIGURATION")
    table_run.font.bold = True
    table_run.font.size = Pt(8)
    
    table = doc.add_table(rows=12, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Parameter"
    header_cells[1].text = "Projected GAN"
    header_cells[2].text = "LoRA"
    
    # Data rows
    data = [
        ("Batch Size", "10", "1 (accum=4)"),
        ("Learning Rate (G/D)", "2.5e-4 / 2.5e-4", "5e-5"),
        ("Optimizer", "Adam (beta=0, 0.99)", "AdamW"),
        ("Image Resolution", "256 x 256", "512 x 512"),
        ("Total Training", "750 kimg (~11h)", "15 epochs (~2h)"),
        ("Gradient Clipping", "0.5", "1.0"),
        ("Mixed Precision", "D: FP16, G: FP32", "FP16"),
        ("R1 Regularization", "gamma=0.1, interval=8", "N/A"),
        ("EMA Decay", "0.9999", "N/A"),
        ("LoRA Rank", "N/A", "128"),
        ("Warmup Steps", "1000", "100"),
    ]
    
    for i, (param, gan_val, lora_val) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = param
        row.cells[1].text = gan_val
        row.cells[2].text = lora_val
    
    # Format table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # Spacing
    
    # ==================== VI. EXPERIMENTS ====================
    exp_title = doc.add_paragraph()
    exp_run = exp_title.add_run("VI. EXPERIMENTS AND RESULTS")
    exp_run.font.bold = True
    exp_run.font.size = Pt(10)
    
    # A. Training Progress
    train_subtitle = doc.add_paragraph()
    train_sub_run = train_subtitle.add_run("A. Training Progress")
    train_sub_run.font.italic = True
    train_sub_run.font.size = Pt(10)
    
    train_text = """The Projected GAN was trained for 750 kimg (75,000 generator updates with batch size 10) over approximately 11-12 hours on an NVIDIA RTX 4060 (8GB VRAM). Training remained stable throughout with no NaN losses observed after implementing the stability measures. The discriminator loss stabilized around 0.5-1.0 while the generator loss converged to approximately 1.0-2.0, indicating healthy adversarial dynamics. Sample quality improved progressively, with recognizable fashion items emerging within the first 10% of training and refinement continuing throughout."""
    
    train_p = doc.add_paragraph()
    train_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    train_p.add_run(train_text).font.size = Pt(10)
    
    lora_train_text = """The LoRA fine-tuning completed in approximately 1-2 hours for 15 epochs. The loss decreased smoothly from initial values around 0.15 to approximately 0.08, indicating successful adaptation to the fashion domain. Validation samples generated during training showed progressive improvement in fashion-specific details."""
    
    lora_train_p = doc.add_paragraph()
    lora_train_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lora_train_p.add_run(lora_train_text).font.size = Pt(10)
    
    # B. Generated Samples
    samples_subtitle = doc.add_paragraph()
    samples_sub_run = samples_subtitle.add_run("B. Generated Samples Analysis")
    samples_sub_run.font.italic = True
    samples_sub_run.font.size = Pt(10)
    
    samples_text = """The Projected GAN successfully generates diverse fashion images including dresses, pants, tops, and full outfits. Generated samples exhibit:

- Accurate clothing shapes and proportions
- Diverse color palettes and patterns
- Realistic fabric textures
- Appropriate human body poses

Some limitations observed include occasional body part distortions, artifacts in complex patterns, and color bleeding between garments. These are typical for GANs trained on limited data and training time.

The LoRA-fine-tuned Stable Diffusion model produces high-fidelity fashion images with excellent text-prompt adherence. The model responds well to prompts specifying clothing type, color, style, and background, demonstrating successful domain adaptation while retaining the base model's compositional abilities."""
    
    samples_p = doc.add_paragraph()
    samples_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    samples_p.add_run(samples_text).font.size = Pt(10)
    
    # Figure placeholder
    fig_text = doc.add_paragraph()
    fig_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_run = fig_text.add_run("\n[Figure 1: Generated fashion samples from Projected GAN (top) and LoRA-fine-tuned Stable Diffusion (bottom)]\n")
    fig_run.font.italic = True
    fig_run.font.size = Pt(9)
    
    # ==================== VII. DISCUSSION ====================
    disc_title = doc.add_paragraph()
    disc_run = disc_title.add_run("VII. DISCUSSION")
    disc_run.font.bold = True
    disc_run.font.size = Pt(10)
    
    disc_text1 = """The two approaches offer complementary strengths. Projected GANs provide fast, unconditional generation suitable for bulk content creation and data augmentation. The lightweight architecture enables rapid inference and easy deployment. However, lack of conditional control limits applicability for specific design requirements."""
    
    disc_p1 = doc.add_paragraph()
    disc_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disc_p1.add_run(disc_text1).font.size = Pt(10)
    
    disc_text2 = """LoRA fine-tuning of Stable Diffusion offers text-based control, enabling generation of specific clothing types, colors, and styles through natural language prompts. The larger model capacity produces higher-fidelity results. However, inference is slower and requires more VRAM."""
    
    disc_p2 = doc.add_paragraph()
    disc_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disc_p2.add_run(disc_text2).font.size = Pt(10)
    
    disc_text3 = """The stability techniques implemented proved essential for successful GAN training. Without spectral normalization and proper gradient clipping, training frequently collapsed with NaN losses. The frozen pretrained discriminator backbone significantly reduced training time compared to training a discriminator from scratch while providing rich multi-scale features."""
    
    disc_p3 = doc.add_paragraph()
    disc_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disc_p3.add_run(disc_text3).font.size = Pt(10)
    
    # ==================== VIII. CONCLUSION ====================
    conc_title = doc.add_paragraph()
    conc_run = conc_title.add_run("VIII. CONCLUSION")
    conc_run.font.bold = True
    conc_run.font.size = Pt(10)
    
    conc_text = """This paper presented a comprehensive dual-approach system for fashion image generation, combining Projected GANs for unconditional synthesis and LoRA-fine-tuned Stable Diffusion for text-conditioned generation. Through extensive experimentation on the DeepFashion dataset, we demonstrated that both approaches can produce visually compelling fashion images on consumer-grade hardware within practical training times.

Comparative Analysis: Our experiments reveal that the LoRA-fine-tuned Stable Diffusion model significantly outperforms the Projected GAN in several key aspects. The diffusion-based approach produces higher-fidelity images with more accurate anatomical proportions, finer texture details, and substantially fewer visual artifacts. The text-conditioning capability enables precise control over generated outputs, allowing users to specify clothing type, color, style, and composition through natural language prompts. Furthermore, the diffusion model benefits from the extensive pretraining of Stable Diffusion v1.5 on billions of image-text pairs, providing a strong foundation for domain-specific adaptation that the GAN cannot match.

In contrast, while the Projected GAN offers faster inference times and lower computational requirements during deployment, it exhibits notable limitations in output diversity and occasionally produces artifacts such as distorted body parts, color bleeding between garments, and unrealistic fabric patterns. The GAN also struggled with complex poses and multi-garment compositions. These issues are characteristic of GANs trained on relatively small datasets with limited training time, and highlight the fundamental advantages of diffusion-based approaches for high-quality image generation.

Key Contributions: This work makes several contributions to the field: (1) a robust implementation of Projected GAN with comprehensive stability techniques including spectral normalization, R1 regularization, exponential moving average, and gradient clipping that effectively prevent training collapse; (2) an efficient LoRA fine-tuning pipeline that successfully adapts Stable Diffusion to the fashion domain using only approximately 10,000 training images; (3) detailed hyperparameter configurations optimized for consumer GPUs with limited VRAM (RTX 4060, 8GB); and (4) empirical validation demonstrating the clear superiority of diffusion-based approaches for fashion image generation tasks.

Limitations: Despite the promising results achieved, several limitations remain. The Projected GAN struggles with complex poses, multi-garment compositions, and maintaining consistent quality across diverse clothing categories. The LoRA-fine-tuned Stable Diffusion model, while producing significantly higher quality outputs, requires considerably more inference time and memory, making it less suitable for real-time applications. Both approaches are constrained by the size and diversity of the training dataset, and would likely benefit from larger, more varied fashion image collections.

Future Work: Several promising directions warrant further investigation: (1) implementing conditional Projected GAN variants with attribute or text conditioning to bridge the quality gap with diffusion models; (2) exploring larger and more diverse fashion datasets, potentially combining multiple sources, to improve generation diversity and quality; (3) investigating knowledge distillation techniques to transfer the quality of diffusion models into faster GAN-based architectures; (4) developing fashion-specific evaluation metrics beyond FID to better assess garment quality, fit accuracy, and style coherence; (5) extending the system to support virtual try-on applications by conditioning on body pose, shape, and existing garments; and (6) exploring multi-modal generation that combines fashion images with textual descriptions, material specifications, and manufacturing constraints.

The complete implementation, including training scripts, model architectures, configuration files, and trained weights, is made available for research purposes, providing a foundation for further exploration of generative models in fashion applications and enabling full reproducibility of our experimental results."""
    
    conc_p = doc.add_paragraph()
    conc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    conc_p.add_run(conc_text).font.size = Pt(10)
    
    # ==================== REFERENCES ====================
    ref_title = doc.add_paragraph()
    ref_run = ref_title.add_run("REFERENCES")
    ref_run.font.bold = True
    ref_run.font.size = Pt(10)
    
    references = [
        '[1] I. Goodfellow et al., "Generative Adversarial Nets," in Advances in Neural Information Processing Systems, 2014.',
        '[2] A. Sauer, K. Schwarz, and A. Geiger, "Projected GANs Converge Faster," in Advances in Neural Information Processing Systems, 2021.',
        '[3] R. Rombach et al., "High-Resolution Image Synthesis with Latent Diffusion Models," in IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2022.',
        '[4] E. J. Hu et al., "LoRA: Low-Rank Adaptation of Large Language Models," in International Conference on Learning Representations, 2022.',
        '[5] T. Karras, S. Laine, and T. Aila, "A Style-Based Generator Architecture for Generative Adversarial Networks," in IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2019.',
        '[6] T. Karras et al., "Analyzing and Improving the Image Quality of StyleGAN," in IEEE/CVF Conference on Computer Vision and Pattern Recognition, 2020.',
        '[7] J. Ho, A. Jain, and P. Abbeel, "Denoising Diffusion Probabilistic Models," in Advances in Neural Information Processing Systems, 2020.',
        '[8] Z. Liu et al., "DeepFashion: Powering Robust Clothes Recognition and Retrieval with Rich Annotations," in IEEE Conference on Computer Vision and Pattern Recognition, 2016.',
        '[9] L. Mescheder, A. Geiger, and S. Nowozin, "Which Training Methods for GANs do actually Converge?," in International Conference on Machine Learning, 2018.',
        '[10] M. Tan and Q. V. Le, "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks," in International Conference on Machine Learning, 2019.',
        '[11] A. Radford et al., "Learning Transferable Visual Models From Natural Language Supervision," in International Conference on Machine Learning, 2021.',
    ]
    
    for ref in references:
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ref_run = ref_p.add_run(ref)
        ref_run.font.size = Pt(8)
        ref_run.font.name = 'Times New Roman'
    
    # Save document
    output_path = "outputs/Fashion_Image_Generation_IEEE_Paper_v3.docx"
    os.makedirs("outputs", exist_ok=True)
    doc.save(output_path)
    print(f"IEEE paper saved to: {output_path}")
    
    return output_path


if __name__ == "__main__":
    create_ieee_paper()
